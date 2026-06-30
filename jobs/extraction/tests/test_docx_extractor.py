"""DocxExtractor — real python-docx text extraction + mocked LLM."""
import json
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from extractors.docx import DocxExtractor, _extract_docx_text


def _make_docx(path: Path):
    d = Document()
    d.add_paragraph("ΤΙΜΟΛΟΓΙΟ ΠΩΛΗΣΗΣ")
    d.add_paragraph("")   # blank — should be skipped
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Σύνολο"
    table.rows[0].cells[1].text = "1500"
    d.save(str(path))


def test_extract_docx_text_includes_paragraphs_and_table(tmp_path):
    p = tmp_path / "doc.docx"
    _make_docx(p)
    text = _extract_docx_text(p)
    assert "ΤΙΜΟΛΟΓΙΟ ΠΩΛΗΣΗΣ" in text
    assert "Σύνολο" in text and "1500" in text


def test_can_handle():
    ex = DocxExtractor()
    assert ex.can_handle(Path("a.docx")) is True
    assert ex.can_handle(Path("a.doc")) is True
    assert ex.can_handle(Path("a.pdf")) is False


def _fake_client(content):
    resp = SimpleNamespace(choices=[SimpleNamespace(message=SimpleNamespace(content=content))])
    return SimpleNamespace(chat=SimpleNamespace(completions=SimpleNamespace(create=lambda **k: resp)))


def test_extract_uses_text_then_llm(tmp_path):
    p = tmp_path / "doc.docx"
    _make_docx(p)
    ex = DocxExtractor()
    ex.client = _fake_client(json.dumps({
        "doc_type": "sales", "detected_language": "el",
        "total_amount": 1500, "vendor_name": "My Co",
    }))
    result = ex.extract(p)
    assert result.doc_type.value == "sales"
    assert result.total_amount == 1500.0
    assert result.source_file == "doc.docx"
    assert "ΤΙΜΟΛΟΓΙΟ" in result.raw_text_excerpt
